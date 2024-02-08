
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import os
from docx.oxml.ns import qn
class FileWriter():
    def __init__(self, path:str) -> None:
        self.file_path = path
        self.document = Document()
        self.document.styles['Normal'].font.name = u'宋体'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def write_t1(self, title:str):
        self.document.add_heading(title,level = 1)

    def write_t2(self, title:str):
        self.document.add_heading(title,level = 2)
    

    def write(self, content:str, level = 0):
        if level:
            self.document.add_heading(content,level = level)
        else:
            self.document.add_paragraph(content)

    def write_with_structure(self, title:str, content: list):
        # style = self.document.styles.add_style("宋体", WD_STYLE_TYPE.PARAGRAPH)
        self.document.add_heading(title,level = 3)
        for item in content:
            self.document.add_paragraph(item)
        self.document.add_paragraph('')
        self.document.add_paragraph('')

    def save(self):
        self.document.save(self.file_path)

    def read(self):
        if not os.path.exists(self.file_path):
            assert ValueError("File path not exist!")
        self.document = Document(self.file_path)
        paragraphs = self.document.paragraphs
        # print(paragraphs[1].text)
        # title:str
        # content = []
        # for idx in range(0, paragraphs.__len__()):
        #     if paragraphs[idx].style.name == "Heading 2":
        #         title = paragraphs[idx].text
        #     elif paragraphs[idx].style.name == "Normal" and not paragraphs[2].text == '':
        #         content.append(paragraphs[idx].text)
        for item in paragraphs:
            yield item


if __name__ == "__main__":
    # pass
    # document = Document()
    # document.add_heading('title 1',level = 1)
    # document.add_paragraph('hello')
    # document.add_paragraph('')
    # document.add_heading('title 1',level = 1)
    # document.add_paragraph('hello')
    # document.save("./1.docx")
    pass